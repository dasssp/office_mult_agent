import asyncio
import csv
import hashlib
import json
from html import escape
from io import StringIO
from pathlib import Path
from typing import Protocol
from uuid import uuid4

from app.schemas import RequestContext
from app.schemas.workflows import DataAnalysisResult


class ArtifactService:
    def __init__(
        self,
        root: Path = Path("artifacts"),
        repository: "ArtifactRepository | None" = None,
    ) -> None:
        self.root = root
        self.repository = repository

    async def export_analysis(
        self, result: DataAnalysisResult, context: RequestContext
    ) -> dict[str, str]:
        tenant_dir = self.root / hashlib.sha256(
            context.tenant_id.encode("utf-8")
        ).hexdigest()[:24]
        await asyncio.to_thread(tenant_dir.mkdir, parents=True, exist_ok=True)
        artifact_id = str(uuid4())
        csv_path = tenant_dir / f"{artifact_id}.csv"
        svg_path = tenant_dir / f"{artifact_id}.svg"
        markdown_path = tenant_dir / f"{artifact_id}.md"
        html_path = tenant_dir / f"{artifact_id}.html"
        xlsx_path = tenant_dir / f"{artifact_id}.xlsx"
        docx_path = tenant_dir / f"{artifact_id}.docx"
        pdf_path = tenant_dir / f"{artifact_id}.pdf"
        metadata_path = tenant_dir / f"{artifact_id}.json"
        buffer = StringIO()
        writer = csv.writer(buffer)
        writer.writerow(["column", "null_count"])
        writer.writerows(result.null_counts.items())
        await asyncio.to_thread(csv_path.write_text, buffer.getvalue(), encoding="utf-8")
        width = 480
        bars = "".join(
            f'<text x="10" y="{40 + index * 32}">{escape(name)}: {count}</text>'
            f'<rect x="180" y="{24 + index * 32}" width="{min(count, 1000) * 30}" '
            'height="16" fill="#2563eb" />'
            for index, (name, count) in enumerate(result.null_counts.items())
        )
        svg = (
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" '
            f'height="{max(80, 40 + len(result.null_counts) * 32)}">'
            f'<text x="10" y="18">空值数量</text>{bars}</svg>'
        )
        await asyncio.to_thread(svg_path.write_text, svg, encoding="utf-8")
        markdown = self._markdown_report(result)
        html = (
            "<!doctype html><html><head><meta charset=\"utf-8\"><title>数据分析报告</title>"
            "</head><body><pre>"
            f"{escape(markdown)}</pre></body></html>"
        )
        await asyncio.to_thread(markdown_path.write_text, markdown, encoding="utf-8")
        await asyncio.to_thread(html_path.write_text, html, encoding="utf-8")
        await asyncio.to_thread(self._write_xlsx, xlsx_path, result)
        await asyncio.to_thread(self._write_docx, docx_path, result)
        await asyncio.to_thread(self._write_pdf, pdf_path, result)
        trace = {
            "artifact_id": artifact_id,
            "source_file_id": result.source_file_id,
            "columns": result.columns,
            "analysis_spec": result.analysis_spec,
            "chart": {
                "type": "bar",
                "metric": "null_count",
                "aggregation": "count",
            },
        }
        await asyncio.to_thread(
            metadata_path.write_text,
            json.dumps(trace, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        if self.repository is not None:
            await self.repository.save(
                artifact_id=artifact_id,
                kind="analysis_bundle",
                object_ref=str(tenant_dir),
                sha256=hashlib.sha256(
                    (buffer.getvalue() + svg).encode("utf-8")
                ).hexdigest(),
                context=context,
            )
        return {
            "artifact_id": artifact_id,
            "summary_csv": str(csv_path),
            "chart_svg": str(svg_path),
            "report_markdown": str(markdown_path),
            "report_html": str(html_path),
            "report_excel": str(xlsx_path),
            "report_word": str(docx_path),
            "report_pdf": str(pdf_path),
            "trace_metadata": str(metadata_path),
        }

    @staticmethod
    def _markdown_report(result: DataAnalysisResult) -> str:
        lines = [
            "# 数据分析报告",
            "",
            f"- 数据行数：{result.row_count}",
            f"- 字段数量：{len(result.columns)}",
            f"- 重复行数：{result.duplicate_rows}",
            "",
            "## 数据质量",
        ]
        lines.extend(
            f"- {column}：空值 {count}"
            for column, count in result.null_counts.items()
        )
        if result.quality_warnings:
            lines.extend(["", "## 风险提示", *[f"- {item}" for item in result.quality_warnings]])
        return "\n".join(lines)

    @staticmethod
    def _write_xlsx(path: Path, result: DataAnalysisResult) -> None:
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "数据质量"
        sheet.append(["字段", "空值数量"])
        for item in result.null_counts.items():
            sheet.append(list(item))
        workbook.save(path)

    @staticmethod
    def _write_docx(path: Path, result: DataAnalysisResult) -> None:
        from docx import Document

        document = Document()
        document.add_heading("数据分析报告", level=1)
        document.add_paragraph(f"数据行数：{result.row_count}")
        for column, count in result.null_counts.items():
            document.add_paragraph(f"{column}：空值 {count}")
        document.save(str(path))

    @staticmethod
    def _write_pdf(path: Path, result: DataAnalysisResult) -> None:
        from reportlab.pdfgen.canvas import Canvas

        canvas = Canvas(str(path))
        canvas.drawString(72, 800, "Data Analysis Report")
        canvas.drawString(72, 780, f"Rows: {result.row_count}")
        for index, (column, count) in enumerate(result.null_counts.items()):
            canvas.drawString(72, 760 - index * 18, f"{column}: nulls={count}")
        canvas.save()


class ArtifactRepository(Protocol):
    async def save(
        self,
        *,
        artifact_id: str,
        kind: str,
        object_ref: str,
        sha256: str,
        context: RequestContext,
    ) -> None: ...
